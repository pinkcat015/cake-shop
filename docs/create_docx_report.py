import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level, color):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Arial'
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(20)
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(12)
        heading.paragraph_format.keep_with_next = True
    elif level == 2:
        run.font.size = Pt(15)
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(8)
        heading.paragraph_format.keep_with_next = True
    elif level == 3:
        run.font.size = Pt(12)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        heading.paragraph_format.keep_with_next = True
    return heading

def main():
    doc = docx.Document()

    # Set standard margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Colors
    color_burgundy = RGBColor(107, 17, 17) # #6b1111
    color_dark_grey = RGBColor(60, 60, 60)
    color_gold = RGBColor(184, 154, 91)    # #b89a5b

    # Set normal style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = color_dark_grey

    # ==================== 1. COVER PAGE ====================
    for _ in range(5):
        doc.add_paragraph()

    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_univ = p_univ.add_run("TRƯỜNG ĐẠI HỌC BÁCH KHOA / CÔNG NGHỆ THÔNG TIN\nKHOA KHOA HỌC & KỸ THUẬT MÁY TÍNH")
    run_univ.bold = True
    run_univ.font.size = Pt(13)
    run_univ.font.color.rgb = color_dark_grey

    for _ in range(4):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("ĐỒ ÁN TỐT NGHIỆP / BÁO CÁO CỰC CHI TIẾT\n\nXÂY DỰNG WEBSITE THƯƠNG MẠI ĐIỆN TỬ\nSCARLETT CAKE SHOP TÍCH HỢP\nCỔNG VIETQR REALTIME SIMULATOR")
    run_title.bold = True
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = color_burgundy

    for _ in range(5):
        doc.add_paragraph()

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta.paragraph_format.left_indent = Inches(1.5)
    
    r_meta = p_meta.add_run(
        "Sinh viên thực hiện:  Nguyễn Minh Trang\n"
        "Mã số sinh viên:     SV999999\n"
        "Ngành học:           Kỹ thuật Phần mềm\n"
        "Lớp:                 KHTT-K18\n"
        "Giảng viên hướng dẫn: PGS. TS. Nguyễn Văn A"
    )
    r_meta.font.size = Pt(12)
    r_meta.line_spacing = 1.3

    for _ in range(5):
        doc.add_paragraph()

    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_footer = p_footer.add_run("HÀ NỘI, THÁNG 6 NĂM 2026")
    r_footer.bold = True
    r_footer.font.size = Pt(11)

    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    add_heading_styled(doc, "MỤC LỤC CHI TIẾT", 1, color_burgundy)
    doc.add_paragraph("LỜI MỞ ĐẦU ............................................................................................................... 3")
    doc.add_paragraph("CHƯƠNG 1: GIỚI THIỆU ĐỀ TÀI VÀ CÔNG NGHỆ SỬ DỤNG ........................................... 4")
    doc.add_paragraph("   1.1. Lý do chọn đề tài .................................................................................................. 4")
    doc.add_paragraph("   1.2. Công nghệ triển khai hệ thống ................................................................................ 5")
    doc.add_paragraph("CHƯƠNG 2: PHÂN TÍCH YÊU CẦU HỆ THỐNG (REQUIREMENTS ANALYSIS) ......... 8")
    doc.add_paragraph("   2.1. Phân tích yêu cầu chức năng (Actors & Usecases) .................................................. 8")
    doc.add_paragraph("   2.2. Phân tích yêu cầu phi chức năng (Non-functional) .................................................. 11")
    doc.add_paragraph("CHƯƠNG 3: THIẾT KẾ CƠ SỞ DỮ LIỆU VÀ KIẾN TRÚC MÃ NGUỒN ............................. 13")
    doc.add_paragraph("   3.1. Thiết kế Mô hình Thực thể Quan hệ (ERD) ............................................................. 13")
    doc.add_paragraph("   3.2. Đặc tả cấu trúc các bảng cơ sở dữ liệu ..................................................................... 14")
    doc.add_paragraph("CHƯƠNG 4: ĐẶC TẢ API ENDPOINTS VÀ BIÊN DỊCH DỮ LIỆU .................................... 19")
    doc.add_paragraph("   4.1. Phân hệ API Hệ thống .................................──────────────────......... 19")
    doc.add_paragraph("CHƯƠNG 5: CÁC LUỒNG XỬ LÝ CỐT LÕI VÀ KỸ THUẬT TỐI ƯU HÓA ......................... 24")
    doc.add_paragraph("   5.1. Luồng Xác thực Email & Mailer Fallback ................................................................. 24")
    doc.add_paragraph("   5.2. Cổng Thanh toán VietQR & Đối soát Realtime ............................................................ 26")
    doc.add_paragraph("   5.3. Tối ưu Bản đồ Định vị GPS Cửa hàng và Chỉ đường Fallback ...................................... 29")
    doc.add_paragraph("CHƯƠNG 6: PHÂN QUYỀN BẢO MẬT HỆ THỐNG .......................................................... 33")
    doc.add_paragraph("CHƯƠNG 7: KẾT QUẢ KIỂM THỬ VÀ ĐÁNH GIÁ ĐỒ ÁN ................................................. 37")
    doc.add_paragraph("KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN .............................................................................. 45")
    doc.add_page_break()

    # ==================== LỜI MỞ ĐẦU ====================
    add_heading_styled(doc, "LỜI MỞ ĐẦU", 1, color_burgundy)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(10)
    p.add_run(
        "Thương mại điện tử đang chứng kiến những bước phát triển vượt bậc tại Việt Nam, đặc biệt là trong lĩnh vực "
        "F&B (Thực phẩm và Đồ uống). Khách hàng ngày nay yêu cầu sự tiện lợi tối đa khi đặt hàng qua mạng, từ việc "
        "lựa chọn sản phẩm, định vị cửa hàng giao bánh tươi nóng hổi, cho đến khâu thanh toán không dùng tiền mặt. "
        "Mã QR ngân hàng theo tiêu chuẩn VietQR Napas247 ra đời đã tạo nên một cuộc cách mạng lớn, cho phép giao dịch chuyển khoản "
        "diễn ra tức thì chỉ với một hành động quét camera điện thoại. Đối với các cửa hàng F&B như Scarlett Cake Shop, "
        "việc tự động hóa khâu xác nhận tiền gửi từ khách hàng đóng vai trò sống còn giúp tối ưu nhân sự và đẩy nhanh tốc độ làm bánh."
    )

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1.25
    p2.paragraph_format.space_after = Pt(10)
    p2.add_run(
        "Mục tiêu cốt lõi của đồ án này là xây dựng một trang thương mại điện tử chuyên nghiệp cung cấp đầy đủ chức năng quản trị và "
        "mua sắm bánh ngọt trực tuyến. Đồ án tập trung nghiên cứu và phát triển giải pháp tích hợp cổng thanh toán chuyển khoản "
        "VietQR động, tự động điền thông tin đơn hàng và tiền phải trả. Song song đó, hệ thống tích hợp luồng đối soát chuyển khoản "
        "giả lập thời gian thực (realtime) giúp đồng bộ ngay lập tức trạng thái giao dịch giữa người mua, nhân viên và người quản lý. "
        "Đây là nền tảng kỹ thuật vững chắc để doanh nghiệp tiến tới tự động hóa quy trình vận hành và kinh doanh trực tuyến."
    )

    doc.add_page_break()

    # ==================== CHƯƠNG 1 ====================
    add_heading_styled(doc, "CHƯƠNG 1: GIỚI THIỆU ĐỀ TÀI VÀ CÔNG NGHỆ SỬ DỤNG", 1, color_burgundy)
    
    add_heading_styled(doc, "1.1. Lý do chọn đề tài", 2, color_burgundy)
    p = doc.add_paragraph("Hiện nay, hầu hết các cửa hàng bán lẻ truyền thống gặp nhiều khó khăn trong khâu quản lý doanh số bán hàng online và thủ tục kiểm soát giao dịch chuyển khoản thủ công. Nhân viên phải thường xuyên kiểm tra điện thoại biến động số dư hoặc yêu cầu khách hàng chụp ảnh màn hình giao dịch (bill chuyển tiền), dẫn đến nguy cơ làm giả hóa đơn hoặc chậm trễ tiến độ hoàn thành đơn bánh tươi. Để khắc phục triệt để lỗ hổng này, Scarlett Cake Shop được thiết kế nhằm tự động hóa toàn bộ quy trình mua hàng - thanh toán - đối soát đơn hàng trực tuyến.")

    add_heading_styled(doc, "1.2. Công nghệ triển khai hệ thống", 2, color_burgundy)
    doc.add_paragraph("Hệ thống sử dụng các công nghệ Web hiện đại bao gồm:")
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("ReactJS (Vite): ").bold = True
    bp1.add_run("Thư viện xây dựng giao diện dựa trên cơ chế Single Page Application (SPA), giúp tối ưu hóa thời gian tải trang và đem lại cảm giác mượt mà cho khách hàng khi thao tác giỏ hàng và xem bản đồ chi nhánh.")

    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.add_run("Express.js (Node.js backend): ").bold = True
    bp2.add_run("Framework xây dựng API RESTful tốc độ cao, khả năng xử lý bất đồng bộ (non-blocking I/O) tốt thích hợp cho việc xử lý các kết nối đối soát đồng thời.")

    bp3 = doc.add_paragraph(style='List Bullet')
    bp3.add_run("MySQL Database: ").bold = True
    bp3.add_run("Hệ quản trị cơ sở dữ liệu quan hệ lưu giữ thông tin sản phẩm, đơn hàng, hóa đơn giao dịch và tài khoản người dùng an toàn, toàn vẹn dữ liệu.")

    bp4 = doc.add_paragraph(style='List Bullet')
    bp4.add_run("Nodemailer SMTP Service: ").bold = True
    bp4.add_run("Module chuyên dụng gửi thư điện tử kích hoạt tài khoản và OTP bảo mật khôi phục mật khẩu thông qua SMTP chính thức của Google Gmail.")

    bp5 = doc.add_paragraph(style='List Bullet')
    bp5.add_run("VietQR API Gateway: ").bold = True
    bp5.add_run("Tự động kết xuất mã QR động dựa trên số tiền đơn hàng và mã đơn, giảm thiểu thao tác nhập tay sai lệch thông tin của khách hàng.")

    doc.add_page_break()

    # ==================== CHƯƠNG 2 ====================
    add_heading_styled(doc, "CHƯƠNG 2: PHÂN TÍCH YÊU CẦU HỆ THỐNG", 1, color_burgundy)
    
    add_heading_styled(doc, "2.1. Phân tích yêu cầu chức năng (Usecase)", 2, color_burgundy)
    doc.add_paragraph("Hệ thống phục vụ ba nhóm vai trò người dùng chính:")

    doc.add_paragraph("1. Khách hàng (Customer):", style='Normal').runs[0].bold = True
    doc.add_paragraph("Khách hàng có thể đăng ký tài khoản, kích hoạt qua link email, đăng nhập và khôi phục mật khẩu qua mã OTP gửi về hòm thư điện tử. Trên ứng dụng, khách hàng được duyệt menu bánh ngọt đa dạng phân loại, thêm bánh vào giỏ hàng, áp dụng các mã voucher giảm giá và tiến hành đặt hàng. Hệ thống hỗ trợ định vị vị trí người dùng để tìm chi nhánh cửa hàng Scarlett gần nhất trong bán kính địa lý. Khi chọn hình thức chuyển khoản, khách hàng quét mã VietQR và xác nhận giả lập chuyển khoản để được chuyển trạng thái đơn hàng ngay lập tức.")

    doc.add_paragraph("2. Nhân viên cửa hàng (Employee):", style='Normal').runs[0].bold = True
    doc.add_paragraph("Nhân viên đại diện cho các chi nhánh quản lý đơn. Nhân viên có thể xem chi tiết tất cả các đơn hàng trong hệ thống, lọc theo mã và cập nhật trạng thái đơn (Chờ xử lý, Đã xác nhận, Đang giao, Đã giao, Hủy đơn). Nhân viên cũng có thể cập nhật thông tin sản phẩm bánh ngọt (tên, giá cả, mô tả, ảnh bánh) tại chi nhánh nhưng không có quyền xóa sản phẩm nhằm đảm bảo tính an toàn dữ liệu.")

    doc.add_paragraph("3. Quản trị viên tối cao (Admin):", style='Normal').runs[0].bold = True
    doc.add_paragraph("Admin nắm toàn bộ quyền kiểm soát hệ thống. Admin truy cập Dashboard đồ họa biểu đồ SVG báo cáo doanh số, số lượng đơn hàng đặt, lượng khách hàng và top bánh bán chạy. Admin có quyền thực hiện đầy đủ thao tác thêm, sửa, xóa bánh ngọt, cấu hình CRUD hệ thống Voucher giảm giá và danh mục các chi nhánh cửa hàng.")

    add_heading_styled(doc, "2.2. Phân tích yêu cầu phi chức năng", 2, color_burgundy)
    doc.add_paragraph("Để triển khai hệ thống ổn định, các yêu cầu phi chức năng sau được đáp ứng:")
    doc.add_paragraph("- Bảo mật: Mật khẩu khách hàng được mã hóa băm một chiều bằng Bcrypt trước khi lưu vào DB. Phiên làm việc của khách hàng bảo mật bằng chuỗi mã hóa ký số JWT Token có thời hạn sử dụng.")
    doc.add_paragraph("- Tính khả dụng: Giao diện thiết kế tương thích mọi độ phân giải màn hình (Responsive Design) từ máy tính bàn, máy tính bảng đến điện thoại di động.")
    doc.add_paragraph("- Tính toàn vẹn: Ràng buộc khóa ngoại trong MySQL đảm bảo khi xóa chi nhánh hoặc bánh, các đơn hàng lịch sử vẫn giữ nguyên vẹn thông tin doanh thu đối soát.")

    doc.add_page_break()

    # ==================== CHƯƠNG 3 ====================
    add_heading_styled(doc, "CHƯƠNG 3: THIẾT KẾ CƠ SỞ DỮ LIỆU", 1, color_burgundy)
    
    add_heading_styled(doc, "3.1. Đặc tả chi tiết các bảng cơ sở dữ liệu", 2, color_burgundy)
    doc.add_paragraph("Dưới đây là đặc tả chi tiết của các bảng dữ liệu cốt lõi trong MySQL:")

    # Table 1: User
    add_heading_styled(doc, "Bảng User (Tài khoản người dùng)", 3, color_burgundy)
    table_user = doc.add_table(rows=5, cols=4)
    table_user.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"]
    for i, title_text in enumerate(headers):
        cell = table_user.rows[0].cells[i]
        cell.text = title_text
        set_cell_background(cell, "6B1111")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True
        set_cell_margins(cell)

    user_rows = [
        ["user_id", "INT", "PRIMARY KEY, AUTO_INCREMENT", "Mã định danh duy nhất"],
        ["username", "VARCHAR(50)", "UNIQUE, NOT NULL", "Tên tài khoản đăng nhập"],
        ["email", "VARCHAR(100)", "UNIQUE, NOT NULL", "Email nhận thông báo và xác thực"],
        ["is_verified", "TINYINT(1)", "DEFAULT 0", "Xác thực email (1: Rồi, 0: Chưa)"]
    ]
    for row_idx, data in enumerate(user_rows):
        row_cells = table_user.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx])
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "FAF6F0")

    doc.add_paragraph() # spacing

    # Table 2: Order
    add_heading_styled(doc, "Bảng Order (Đơn đặt hàng)", 3, color_burgundy)
    table_order = doc.add_table(rows=6, cols=4)
    table_order.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, title_text in enumerate(headers):
        cell = table_order.rows[0].cells[i]
        cell.text = title_text
        set_cell_background(cell, "6B1111")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True
        set_cell_margins(cell)

    order_rows = [
        ["order_id", "INT", "PRIMARY KEY, AUTO_INCREMENT", "Mã số đơn đặt hàng"],
        ["total_price", "DECIMAL(10,2)", "NOT NULL", "Tổng số tiền thanh toán của đơn"],
        ["status", "VARCHAR(20)", "DEFAULT 'PENDING'", "Trạng thái đơn hàng hiện tại"],
        ["delivery_method", "VARCHAR(20)", "NOT NULL", "Hình thức giao (delivery / pickup)"],
        ["store_id", "INT", "FOREIGN KEY", "Chi nhánh chịu trách nhiệm giao hàng"]
    ]
    for row_idx, data in enumerate(order_rows):
        row_cells = table_order.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx])
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "FAF6F0")

    doc.add_page_break()

    # ==================== CHƯƠNG 4 ====================
    add_heading_styled(doc, "CHƯƠNG 4: ĐẶC TẢ API ENDPOINTS VÀ BIÊN DỊCH DỮ LIỆU", 1, color_burgundy)
    
    doc.add_paragraph("Các API RESTful chính được kết nối giữa Client React và Server Node.js:")

    table_api = doc.add_table(rows=6, cols=4)
    table_api.alignment = WD_TABLE_ALIGNMENT.CENTER
    api_headers = ["Method", "Endpoint Route", "Phân quyền", "Mô tả chức năng"]
    for i, title_text in enumerate(api_headers):
        cell = table_api.rows[0].cells[i]
        cell.text = title_text
        set_cell_background(cell, "6B1111")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True
        set_cell_margins(cell)

    api_rows = [
        ["POST", "/api/auth/register", "Public", "Đăng ký tài khoản và tạo verification token"],
        ["GET", "/api/auth/verify-email", "Public", "Kích hoạt tài khoản khi bấm link trong email"],
        ["POST", "/api/orders", "Customer", "Tạo đơn đặt hàng mới từ danh sách giỏ hàng"],
        ["POST", "/api/payments/confirm", "Customer", "Đối soát và xác nhận thanh toán đơn hàng thành PAID"],
        ["GET", "/api/reports/stats", "Admin", "Lấy thống kê doanh số tổng hợp vẽ lên dashboard"]
    ]
    for row_idx, data in enumerate(api_rows):
        row_cells = table_api.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx])
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "FAF6F0")

    doc.add_page_break()

    # ==================== CHƯƠNG 5 ====================
    add_heading_styled(doc, "CHƯƠNG 5: CÁC LUỒNG XỬ LÝ CỐT LÕI VÀ KỸ THUẬT TỐI ƯU HÓA", 1, color_burgundy)
    
    add_heading_styled(doc, "5.1. Luồng Xác thực Email & Mailer Fallback", 2, color_burgundy)
    doc.add_paragraph("Để bảo vệ hệ thống khỏi các tài khoản ảo hoặc hòm thư không có thực, khi khách hàng thực hiện đăng ký, backend sẽ tạo ra một Verification Token ngẫu nhiên bằng chuỗi Hex và lưu vào cơ sở dữ liệu kèm trạng thái is_verified = 0. Sau đó Nodemailer sẽ kết nối đến cổng SMTP Gmail của cửa hàng gửi bức thư chứa liên kết xác nhận. Trong trường hợp thông tin tài khoản SMTP trong file cấu hình .env bị sai hoặc bị chặn bảo mật (BadCredentials), hệ thống sẽ kích hoạt Mailer Fallback tự động: in trực tiếp đường link kích hoạt lên màn hình log của console. Nhờ đó nhà phát triển vẫn có thể copy link và hoàn tất việc kích hoạt tài khoản trong môi trường phát triển (development).")

    add_heading_styled(doc, "5.2. Cổng Thanh toán VietQR & Đối soát Realtime", 2, color_burgundy)
    doc.add_paragraph("Khi đơn hàng thanh toán bằng phương thức chuyển khoản, khách hàng được điều hướng tới màn hình cổng thanh toán. Tại đây, hệ thống gửi yêu cầu sinh ảnh QR động tới API của VietQR dựa trên số tài khoản ngân hàng của cửa hàng (Ngân hàng Quân đội MB, STK: 0348582531, Chủ tài khoản: NGUYEN MINH TRANG) cùng số tiền cần trả và nội dung bắt buộc ghi chính xác để đối soát là 'SCARLETT <order_id>'.")
    doc.add_paragraph("Nút 'Giả lập thanh toán' trên giao diện cho phép khách hàng thực hiện giao dịch thử nghiệm. Khi click, React sẽ kích hoạt màn hình Verifying Overlay với hiệu ứng mờ nền và thanh tiến trình chạy qua 4 bước: kết nối ngân hàng, quét giao dịch khớp nội dung, xác nhận số tiền và đồng bộ trạng thái. Sau khi chạy hết 3.2 giây, frontend gửi POST /confirm lên backend. Cơ sở dữ liệu cập nhật trạng thái hóa đơn là PAID, trạng thái đơn là CONFIRMED và trả về kết quả thành công cho người dùng.")

    add_heading_styled(doc, "5.3. Tối ưu Bản đồ Định vị GPS Cửa hàng và Chỉ đường Fallback", 2, color_burgundy)
    doc.add_paragraph("Để mang lại trải nghiệm tiện ích, hệ thống tích hợp bản đồ Google Maps miễn phí thông qua thẻ iframe. Trong trang chi tiết chi nhánh cửa hàng, nếu chi nhánh bị thiếu thông tin kinh độ/vĩ độ trong cơ sở dữ liệu, nút hướng dẫn đường đi sẽ tự động mã hóa chuỗi địa chỉ chữ (address string) của chi nhánh làm điểm đến trên Google Maps, ngăn ngừa tuyệt đối lỗi chỉ đường tới tọa độ null hoặc hiển thị bản đồ trống.")

    doc.add_page_break()

    # ==================== CHƯƠNG 6 & 7 ====================
    add_heading_styled(doc, "CHƯƠNG 6: PHÂN QUYỀN BẢO MẬT HỆ THỐNG", 1, color_burgundy)
    doc.add_paragraph("Hệ thống bảo vệ dữ liệu ở hai lớp độc lập:")
    doc.add_paragraph("1. Bảo vệ API Backend: Sử dụng middleware authenticateToken để trích xuất token gửi kèm trong HTTP Authorization Header, giải mã lấy ID người dùng và vai trò. Middleware tiếp theo authorizeRoles kiểm tra xem vai trò đó có nằm trong danh sách được cấp phép thực hiện hành động hay không. Nếu không, server lập tức phản hồi mã lỗi 403 Forbidden.")
    doc.add_paragraph("2. Bảo vệ Route Frontend: React sử dụng các component Route Guards bọc ngoài các màn hình quản trị. Khi trang web được tải, React kiểm tra token và vai trò lưu trữ trong context. Nếu người dùng chưa đăng nhập, họ sẽ bị chuyển hướng về màn hình đăng nhập. Nếu vai trò không phù hợp (ví dụ nhân viên cố gắng vào trang Dashboard doanh thu của Admin), họ sẽ được điều hướng về trang chủ.")

    add_heading_styled(doc, "CHƯƠNG 7: KẾT QUẢ KIỂM THỬ VÀ ĐÁNH GIÁ ĐỒ ÁN", 1, color_burgundy)
    doc.add_paragraph("Quá trình chạy thử nghiệm cho thấy kết quả hệ thống hoàn toàn đồng bộ, đáp ứng toàn bộ các yêu cầu của hội đồng chuyên môn:")
    doc.add_paragraph("- Biên dịch mã nguồn Frontend (React Vite build): Thành công 100% không phát sinh lỗi cảnh báo.")
    doc.add_paragraph("- Chạy Suite kiểm thử tự động (test_completed_features.js): Vượt qua tất cả các chặng kiểm tra đổi mật khẩu qua mã OTP, phân quyền chặn truy cập Employee, CRUD Voucher.")
    doc.add_paragraph("- Chạy kiểm thử thanh toán VietQR (test_payment_gateway.js): Hoàn tất luồng thanh toán chuyển khoản, cập nhật DB sang trạng thái CONFIRMED/PAID khớp dữ liệu đối soát thực tế.")

    doc.add_page_break()

    # ==================== KẾT LUẬN ====================
    add_heading_styled(doc, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1, color_burgundy)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.add_run(
        "Hệ thống website Scarlett Cake Shop đã giải quyết được các bài toán quản lý đặt đơn và thanh toán tự động thời gian thực "
        "thực tế cho cửa hàng bánh ngọt. Với kiến trúc tách biệt rõ ràng, ứng dụng chạy ổn định và đạt tốc độ phản hồi tối ưu. "
        "Trong thời gian tới, hướng nâng cấp của đề tài sẽ tập trung kết nối trực tiếp cổng thanh toán thông qua Webhook chuyển khoản "
        "ngân hàng thực để hoàn thiện tối đa tính năng đối soát số dư của Scarlett Cake Shop."
    )

    doc.save("C:\\Users\\pinkc\\cake-shop\\docs\\Bao_cao_do_an_Scarlett_Cake_Shop.docx")
    print("Document created successfully!")

if __name__ == "__main__":
    main()
